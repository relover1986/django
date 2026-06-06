#%%
from django.shortcuts import render, redirect
from django.http import HttpResponse
from .models import Question, UserAnswer
import pandas as pd
import random
import arrow
import sqlite3 as sl
from pathlib import Path
from io import BytesIO
from django.http import JsonResponse
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

BASE = Path(__file__).parent.parent

CATEGORY_CONFIG = {
    "爆破": {
        "title": "辽宁捷祥民用爆破三大员培训题库",
        "url_prefix": "/home/custom_quiz",
        "ti_type": "爆破",
    },
    "井工": {
        "title": "辽宁捷祥非煤矿山井工培训题库",
        "url_prefix": "/home/custom_quiz",
        "ti_type": "非煤矿山井工",
    },
    "危装": {
        "title": "辽宁捷祥危险品装卸培训题库",
        "url_prefix": "/home/custom_quiz",
        "ti_type": "危险品装卸",
    },
    "自定义": {
        "title": "自定义答题",
        "url_prefix": "/home/custom_quiz",
        "ti_type": "自定义",
    },
}

def _get_opt_text(options_str: str, answer: str) -> str:
    """提取选项文字，支持 A.xxx / A xxx / Axxx 格式，支持多选 ABC"""
    if not answer:
        return ""
    results = []
    for ch in answer:
        for opt in options_str.split():
            clean = opt.strip()
            if not clean:
                continue
            if clean[0] == ch:
                txt = clean[1:].lstrip(". \t")
                results.append(txt)
                break
    return " | ".join(results)


def _ti_new_core(request, category, template_name="custom_quiz.html"):
    """统一出题/评分逻辑，按 category 区分题库"""
    cfg = CATEGORY_CONFIG.get(category)
    if cfg is None:
        cfg = {
            "title": f"{category}培训题库",
            "url_prefix": "/home/custom_quiz",
            "ti_type": category,
        }
    title = cfg["title"]
    url_prefix = cfg["url_prefix"]
    ti_type = cfg["ti_type"]
    session_key = f"quiz_data_{category}"
    if url_prefix == "/home/custom_quiz":
        reload_url = url_prefix + "_reload" + "?category=" + category
    else:
        reload_url = url_prefix + "_reload"

    date = arrow.now().shift(days=0).format("YYYY-MM-DD")
    username = request.session["info"]["name"]
    ident = request.session["info"]["ident"]

    if request.method == "GET":
        data = Question.objects.filter(category=category).values(
            "tihao", "question_type", "question", "options", "correct_answer"
        )
        df = pd.DataFrame(data)

        answered = UserAnswer.objects.values("tihao").filter(ident=ident, ti_type=ti_type)
        if answered.exists():
            df_answered = pd.DataFrame(answered)
            df = df.merge(df_answered, on="tihao", how="left", indicator=True)
            df = df[df["_merge"] == "left_only"].drop(columns=["_merge"])

        df.columns = ["tihao", "题型", "题目", "选项", "正确答案"]
        n_all = len(df)
        dfs = []
        for qtype, n_max in [("单选题", 5), ("多选题", 5), ("判断题", 10)]:
            sub = df[df["题型"] == qtype]
            n = min(len(sub), n_max)
            if n > 0:
                sub = sub.sample(n=n)
                sub["序号"] = range(1, n + 1)
                sub["题号"] = qtype.replace("题", "") + sub["序号"].astype(str)
                dfs.append(sub)
        df = pd.concat(dfs) if dfs else pd.DataFrame()
        questions = df.to_dict(orient="records")

        request.session[session_key] = questions

        context = {
            "title": f"{title} -- {username} 还有{n_all}题",
            "questions": questions,
        }
        ctx = context
        if template_name == "custom_quiz.html":
            all_cats = list(Question.objects.values_list("category", flat=True).distinct().order_by("category"))
            ctx = {**context, "category": category, "categories": all_cats}
        return render(request, template_name, ctx)

    # POST: scoring
    quiz_data = request.session.get(session_key, [])
    if not quiz_data:
        return redirect(url_prefix)

    df = pd.DataFrame(quiz_data)
    df["题号"] = df["题型"].str.replace("题", "") + df["序号"].astype(str)

    rows = []
    for _, q in df.iterrows():
        tihao = q["题号"]
        answer = request.POST.getlist(tihao)
        answer = "".join(answer)
        rows.append({"题号": tihao, "我的答案": answer, "tihao": q["tihao"],
                      "题目": q["题目"], "选项": q["选项"], "正确答案": q["正确答案"], "题型": q["题型"],
                      "我的答案文字": _get_opt_text(q["选项"], answer),
                      "正确答案文字": _get_opt_text(q["选项"], q["正确答案"])})

    result_df = pd.DataFrame(rows)
    result_df["我的答案"] = result_df["我的答案"].str.replace(r"[^a-zA-Z0-9\u4e00-\u9fa5]", "", regex=True)
    result_df["得分"] = "正确"
    result_df.loc[result_df["我的答案"] != result_df["正确答案"], "得分"] = "错误"

    correct = result_df[result_df["得分"] == "正确"]
    if not correct.empty:
        records = correct[["tihao"]].copy()
        records["ti_type"] = ti_type
        records["date"] = date
        records["ident"] = ident
        objs = []
        for _, r in records.iterrows():
            objs.append(UserAnswer(ti_type=r["ti_type"], tihao=r["tihao"],
                                   date=r["date"], ident=r["ident"]))
        UserAnswer.objects.bulk_create(objs)

    wrong = result_df[result_df["得分"] == "错误"].copy()
    wrong["分值"] = 1
    wrong.loc[wrong["题号"].str.contains("多选"), "分值"] = 2
    total_deduct = wrong["分值"].sum() if not wrong.empty else 0

    wrong_list = wrong.to_dict(orient="records") if not wrong.empty else []
    correct_list = result_df[result_df["得分"] == "正确"].to_dict(orient="records") if not result_df[result_df["得分"] == "正确"].empty else []

    return render(request, "result_new.html", {
        "title": f"扣分 -- {total_deduct}",
        "total": len(result_df),
        "correct": len(correct),
        "wrong": total_deduct,
        "wrong_list": wrong_list,
        "correct_list": correct_list,
        "reload_url": reload_url,
        "continue_url": url_prefix,
    })


def custom_quiz(request):
    category = request.GET.get("category", "爆破")
    return _ti_new_core(request, category, template_name="custom_quiz.html")


def custom_quiz_reload(request):
    category = request.GET.get("category", "爆破")
    session_key = f"quiz_data_{category}"
    if session_key in request.session:
        del request.session[session_key]
    return redirect(f"/home/custom_quiz?category={category}")


def export_docx(request):
    """导出试卷为 DOCX 文件"""
    page = request.GET.get("page", "custom_quiz")

    # 映射 page 到 session key 和 redirect URL
    page_map = {
        "custom_quiz": {"session_key": "quiz_data_自定义", "redirect_url": "/home/custom_quiz"},
        "custom_quiz": {"session_key": "quiz_data_自定义", "redirect_url": "/home/custom_quiz"},
    }
    cfg = page_map.get(page)
    if not cfg:
        return redirect("/home/custom_quiz")

    quiz_data = request.session.get(cfg["session_key"], [])
    if not quiz_data:
        return redirect(cfg["redirect_url"])

    doc = Document()

    # A4 纸张
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)

    # 页眉：宋体 小五，内容"矿业有限公司"
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("矿业有限公司")
    run.font.name = "宋体"
    run.font.size = Pt(9)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")

    def add_run_fangsong(paragraph, text, bold=False, size=Pt(12)):
        """添加仿宋 4号 文字的辅助函数"""
        run = paragraph.add_run(text)
        run.font.name = "仿宋"
        run.font.size = size
        run.bold = bold
        run._element.rPr.rFonts.set(qn('w:eastAsia'), "仿宋")
        return run

    # === 第1页：试卷（题目，不含答案）===
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_fangsong(p_title, "考  试  试  卷", bold=True, size=Pt(16))
    doc.add_paragraph()

    # 按题型分组输出
    def output_questions(data_list, include_answer=False):
        type_order = {"单选题": "📝 单选题", "多选题": "📝 多选题", "判断题": "📝 判断题"}
        by_type = {}
        for q in data_list:
            t = q["题型"]
            if t not in by_type:
                by_type[t] = []
            by_type[t].append(q)

        for qtype, title in type_order.items():
            if qtype not in by_type:
                continue
            # 题型标题
            p_section = doc.add_paragraph()
            add_run_fangsong(p_section, title, bold=True)

            for q in by_type[qtype]:
                # 题号 + 题目
                p_q = doc.add_paragraph()
                q_text = f"第{q['序号']}题. {q['题目']}"
                add_run_fangsong(p_q, q_text, bold=True)

                # 选项
                if qtype == "单选题":
                    # 单选选项水平排列
                    opts = str(q.get("选项", "")).split()
                    p_opts = doc.add_paragraph()
                    p_opts.paragraph_format.space_before = Pt(2)
                    p_opts.paragraph_format.space_after = Pt(2)
                    first = True
                    for opt in opts:
                        clean = opt.strip()
                        if not clean:
                            continue
                        if not first:
                            add_run_fangsong(p_opts, "    ")
                        add_run_fangsong(p_opts, clean)
                        first = False
                elif qtype == "多选题":
                    # 多选选项水平排列
                    opts = str(q.get("选项", "")).split()
                    p_opts = doc.add_paragraph()
                    p_opts.paragraph_format.space_before = Pt(2)
                    p_opts.paragraph_format.space_after = Pt(2)
                    first = True
                    for opt in opts:
                        clean = opt.strip()
                        if not clean:
                            continue
                        if not first:
                            add_run_fangsong(p_opts, "    ")
                        add_run_fangsong(p_opts, clean)
                        first = False

                elif qtype == "判断题":
                    p_opt = doc.add_paragraph()
                    add_run_fangsong(p_opt, "    对      错")

                if include_answer:
                    correct = q.get("正确答案", "")
                    if qtype == "单选题" or qtype == "多选题":
                        add_run_fangsong(p_opts, f"    【答案：{correct}】")
                    elif qtype == "判断题":
                        add_run_fangsong(p_opt, f"    【答案：{correct}】")

    # 输出试卷（不含答案）
    output_questions(quiz_data, include_answer=False)

    # === 分页 + 答案页 ===
    doc.add_page_break()

    p_ans_title = doc.add_paragraph()
    p_ans_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_fangsong(p_ans_title, "参  考  答  案", bold=True, size=Pt(16))
    doc.add_paragraph()

    # 输出答案（内联到选项后同一行）
    output_questions(quiz_data, include_answer=True)

    # 保存到 BytesIO 返回为下载
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f"考试试卷_{page}.docx"
    response = HttpResponse(
        buffer.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response



def import_questions(request):
    """接收 Excel 文件，解析并导入题库"""
    if request.method != "POST":
        return JsonResponse({"success": False, "error": "仅支持 POST 请求"}, status=405)

    file = request.FILES.get("file")
    if not file:
        return JsonResponse({"success": False, "error": "未上传文件"}, status=400)

    name = file.name.lower()
    if not (name.endswith(".xlsx") or name.endswith(".xls")):
        return JsonResponse({"success": False, "error": "仅支持 .xlsx / .xls 文件"}, status=400)

    try:
        import pandas as pd
        df = pd.read_excel(file)
    except Exception as e:
        return JsonResponse({"success": False, "error": f"文件解析失败: {str(e)}"}, status=400)

    required = {"category", "question_type", "question", "options", "correct_answer"}
    actual = set(df.columns)
    missing = required - actual
    if missing:
        return JsonResponse({"success": False, "error": f"缺少字段: {', '.join(sorted(missing))}"}, status=400)

    # 自定义类别：不限值

    valid_types = {"单选题", "多选题", "判断题"}
    bad_types = set(df["question_type"].dropna().unique()) - valid_types
    if bad_types:
        return JsonResponse({"success": False, "error": f"无效的 question_type 值: {', '.join(str(x) for x in bad_types)}，允许: 单选题/多选题/判断题"}, status=400)

    from app01.models import Question
    objs = []
    errors = []
    for i, row in df.iterrows():
        try:
            cat = str(row.get("category", "")).strip()
            qt = str(row.get("question_type", "")).strip()
            ti = str(row.get("tihao", str(i + 1))).strip()
            q = str(row.get("question", "")).strip()
            opts = str(row.get("options", "")).strip()
            ans = str(row.get("correct_answer", "")).strip()
            if not q or not opts or not ans:
                errors.append(f"第{i+2}行: 题目/选项/答案不能为空")
                continue
            objs.append(Question(
                category=cat, question_type=qt, tihao=ti,
                question=q, options=opts, correct_answer=ans
            ))
        except Exception as e:
            errors.append(f"第{i+2}行: {str(e)}")

    if objs:
        Question.objects.bulk_create(objs)

    msg = f"成功导入 {len(objs)} 题"
    if errors:
        msg += f"，{len(errors)} 条失败"
    return JsonResponse({"success": True, "message": msg, "imported": len(objs), "errors": errors[:5]})



def download_example_questions(request):
    """下载示例 Excel 模板"""
    import pandas as pd
    from io import BytesIO
    from django.http import HttpResponse

    data = {
        "category": ["爆破", "井工", "危装"],
        "question_type": ["单选题", "多选题", "判断题"],
        "tihao": ["001", "002", "003"],
        "question": ["爆破作业时，警戒信号发出后应？", "井工矿山通风系统的作用包括？", "危险品装卸作业可以使用手机吗？"],
        "options": [
            "A.立即起爆 B.确认安全后起爆 C.撤离现场 D.检查线路",
            "A.排除有毒气体 B.提供新鲜空气 C.降低粉尘浓度 D.调节温度",
            "A.可以 B.不可以"
        ],
        "correct_answer": ["B", "ABC", "B"],
    }
    df = pd.DataFrame(data)
    buf = BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="题库模板")
    buf.seek(0)
    resp = HttpResponse(buf.getvalue(), content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    resp["Content-Disposition"] = 'attachment; filename="题库导入模板.xlsx"'
    return resp
